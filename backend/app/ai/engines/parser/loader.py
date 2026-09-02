import os
from pathlib import Path
import fitz  # PyMuPDF
import pdfplumber
import docx
from app.core.logger import logger

def load_pdf_text_fitz(file_path: Path) -> str:
    """Extracts text using PyMuPDF (fitz) - fastest method."""
    text = ""
    try:
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text()
        doc.close()
    except Exception as e:
        logger.warning(f"PyMuPDF failed to extract text from {file_path.name}: {e}")
    return text.strip()

def load_pdf_text_pdfplumber(file_path: Path) -> str:
    """Extracts text using pdfplumber - robust backup for complex tables/layouts."""
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        logger.error(f"pdfplumber failed to extract text from {file_path.name}: {e}")
    return text.strip()

def load_docx_text(file_path: Path) -> str:
    """Extracts text from Word documents using python-docx."""
    text = ""
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        text = "\n".join(full_text)
    except Exception as e:
        logger.error(f"docx failed to extract text from {file_path.name}: {e}")
    return text.strip()

def load_document_text(file_path: Path) -> str:
    """
    Unified loader to select correct parser based on file extension.
    Ensures PDF falls back to secondary parser if primary fails.
    """
    ext = file_path.suffix.lower()
    logger.info(f"Loading document text for: {file_path.name} with extension: {ext}")
    
    if ext == ".pdf":
        text = load_pdf_text_fitz(file_path)
        if not text:
            logger.info("PyMuPDF returned empty string. Retrying with pdfplumber...")
            text = load_pdf_text_pdfplumber(file_path)
        return text
    elif ext in [".docx", ".doc"]:
        return load_docx_text(file_path)
    else:
        logger.error(f"Unsupported file format: {ext}")
        raise ValueError(f"Unsupported file format: {ext}")
